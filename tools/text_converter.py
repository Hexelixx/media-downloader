"""
Onglet : convertisseur de formats de texte (txt / md / html / docx / pdf).

N'importe quel format supporté peut être converti vers n'importe quel autre
(y compris le PDF, en lecture comme en écriture).
"""

import os
import threading
import queue
import traceback
from xml.sax.saxutils import escape as _xml_escape

import customtkinter as ctk
from tkinter import filedialog, messagebox

import markdown
import html2text
import docx
import pymupdf
from docx.document import Document as _DocxDocument
from docx.oxml.table import CT_Tbl
from docx.oxml.text.paragraph import CT_P
from docx.table import Table as _DocxTable, _Cell as _DocxCell
from docx.text.paragraph import Paragraph as _DocxParagraph
from reportlab.lib.pagesizes import A4
from reportlab.lib.units import cm
from reportlab.lib.styles import getSampleStyleSheet
from reportlab.platypus import SimpleDocTemplate, Paragraph as _RLParagraph, Spacer

from common import (
    PAD, DEFAULT_OUTPUT_DIR, ensure_dir, enable_file_drop, build_scrollable_body,
    export_log_to_file,
)
from i18n import t

# Formats que l'outil peut PRODUIRE (menu "Convertir vers") -- .doc (ancien format binaire
# Word) n'y figure pas : personne n'a besoin de générer ce format obsolète.
SUPPORTED_EXTENSIONS = (".txt", ".md", ".html", ".htm", ".docx", ".pdf")

# Formats que l'outil peut LIRE en entrée : les précédents + .doc (lecture seule, via
# automatisation Word -- voir _read_doc_text). Utilisé pour le sélecteur de fichiers et
# le glisser-déposer, afin que les .doc ne soient plus "invisibles" pour l'appli.
READABLE_EXTENSIONS = SUPPORTED_EXTENSIONS + (".doc",)

DEFAULT_CONVERTER_OUTPUT_DIR = os.path.join(DEFAULT_OUTPUT_DIR, t("text.output_subdir"))


# --------------------------------------------------------------------- Logique pure ---

def _iter_block_items(parent):
    """Parcourt les paragraphes ET tableaux d'un document/cellule .docx, dans l'ordre
    où ils apparaissent réellement (recette standard python-docx). Sans ça, le texte
    placé dans des tableaux (très courant dans de vrais documents Word) est ignoré."""
    if isinstance(parent, _DocxDocument):
        parent_elm = parent.element.body
    elif isinstance(parent, _DocxCell):
        parent_elm = parent._tc
    else:
        raise ValueError(t("text.unsupported_docx_container", type=repr(type(parent))))

    for child in parent_elm.iterchildren():
        if isinstance(child, CT_P):
            yield _DocxParagraph(child, parent)
        elif isinstance(child, CT_Tbl):
            yield _DocxTable(child, parent)


def _docx_container_text(container) -> str:
    """Texte complet d'un document .docx (ou d'une cellule de tableau, appelé
    récursivement pour les tableaux imbriqués), paragraphes ET tableaux inclus."""
    parts = []
    for block in _iter_block_items(container):
        if isinstance(block, _DocxParagraph):
            if block.text.strip():
                parts.append(block.text)
        elif isinstance(block, _DocxTable):
            for row in block.rows:
                for cell in row.cells:
                    cell_text = _docx_container_text(cell)
                    if cell_text.strip():
                        parts.append(cell_text)
    return "\n\n".join(parts)


def _read_docx_text(input_path: str) -> str:
    document = docx.Document(input_path)
    text = _docx_container_text(document)
    # Filet de sécurité : si la récupération structurée n'a rien donné (structure
    # inhabituelle), on retombe sur la simple liste de paragraphes plutôt que de
    # rendre un fichier vide.
    if not text.strip():
        text = "\n\n".join(p.text for p in document.paragraphs if p.text.strip())
    return text


def _read_pdf_text(input_path: str) -> str:
    document = pymupdf.open(input_path)
    try:
        parts = [(page.get_text() or "").strip() for page in document]
    finally:
        document.close()
    return "\n\n".join(p for p in parts if p)


def _read_doc_text(input_path: str) -> str:
    """Lit un .doc (ancien format binaire Word 97-2003) via automatisation COM de
    Microsoft Word. Aucune bibliothèque Python pure ne sait lire ce format (contrairement
    au .docx, qui est un simple zip) -- Word installé sur la machine est donc requis.
    Ouvre le fichier en lecture seule et ne l'enregistre jamais, ferme Word proprement
    même en cas d'erreur (jamais de processus WINWORD.EXE orphelin)."""
    try:
        import pythoncom
        import win32com.client as win32
    except ImportError as e:
        raise RuntimeError(t("text.doc_needs_word")) from e

    pythoncom.CoInitialize()
    word = None
    doc = None
    try:
        word = win32.Dispatch("Word.Application")
        word.Visible = False
        word.DisplayAlerts = False
        doc = word.Documents.Open(
            os.path.abspath(input_path), ConfirmConversions=False,
            ReadOnly=True, AddToRecentFiles=False,
        )
        return doc.Content.Text
    except Exception as e:
        raise RuntimeError(t("text.doc_read_failed", error=e)) from e
    finally:
        if doc is not None:
            try:
                doc.Close(SaveChanges=False)
            except Exception:
                pass
        if word is not None:
            try:
                word.Quit()
            except Exception:
                pass
        pythoncom.CoUninitialize()


# Codes WdSaveFormat de Word pour l'enregistrement natif (voir _convert_doc_via_word).
_WORD_SAVE_FORMAT = {".docx": 16, ".pdf": 17}  # wdFormatDocumentDefault, wdFormatPDF


def _convert_doc_via_word(input_path: str, output_path: str, target_ext: str) -> None:
    """Convertit un .doc en .docx ou .pdf directement via Word ('Enregistrer sous' /
    export natif), ce qui préserve la mise en page (polices, gras/italique, tableaux,
    images, styles, en-têtes...) -- contrairement au pivot texte brut générique
    (_read_pivot_text/_write_pivot_text) qui réduit tout à de simples paragraphes.
    Nécessite Word installé (voir _read_doc_text pour le message d'erreur sinon)."""
    try:
        import pythoncom
        import win32com.client as win32
    except ImportError as e:
        raise RuntimeError(t("text.doc_convert_needs_word")) from e

    format_code = _WORD_SAVE_FORMAT[target_ext]

    pythoncom.CoInitialize()
    word = None
    doc = None
    try:
        word = win32.Dispatch("Word.Application")
        word.Visible = False
        word.DisplayAlerts = False
        doc = word.Documents.Open(
            os.path.abspath(input_path), ConfirmConversions=False,
            ReadOnly=True, AddToRecentFiles=False,
        )
        out_dir = os.path.dirname(os.path.abspath(output_path))
        if out_dir:
            ensure_dir(out_dir)
        doc.SaveAs2(os.path.abspath(output_path), FileFormat=format_code)
    except Exception as e:
        raise RuntimeError(t("text.doc_convert_failed", ext=target_ext, error=e)) from e
    finally:
        if doc is not None:
            try:
                doc.Close(SaveChanges=False)
            except Exception:
                pass
        if word is not None:
            try:
                word.Quit()
            except Exception:
                pass
        pythoncom.CoUninitialize()


def _read_pivot_text(input_path: str) -> str:
    """Lit un fichier source et retourne un texte "pivot" (texte simple / markdown lisible)."""
    ext = os.path.splitext(input_path)[1].lower()

    if ext in (".txt", ".md"):
        with open(input_path, "r", encoding="utf-8", errors="replace") as f:
            return f.read()

    if ext in (".html", ".htm"):
        with open(input_path, "r", encoding="utf-8", errors="replace") as f:
            content = f.read()
        return html2text.html2text(content)

    if ext == ".docx":
        return _read_docx_text(input_path)

    if ext == ".doc":
        return _read_doc_text(input_path)

    if ext == ".pdf":
        return _read_pdf_text(input_path)

    raise ValueError(t("text.unsupported_source", ext=ext, accepted=", ".join(READABLE_EXTENSIONS)))


def _write_pdf(pivot_text: str, output_path: str) -> None:
    out_dir = os.path.dirname(os.path.abspath(output_path))
    if out_dir:
        ensure_dir(out_dir)

    styles = getSampleStyleSheet()
    style = styles["Normal"]

    doc = SimpleDocTemplate(
        output_path, pagesize=A4,
        leftMargin=2 * cm, rightMargin=2 * cm, topMargin=2 * cm, bottomMargin=2 * cm,
    )

    blocks = [b for b in pivot_text.strip().split("\n\n") if b.strip()]
    story = []
    if not blocks:
        story.append(_RLParagraph("", style))
    else:
        for block in blocks:
            safe = _xml_escape(block).replace("\r\n", "\n").replace("\n", "<br/>")
            story.append(_RLParagraph(safe, style))
            story.append(Spacer(1, 8))
    doc.build(story)


def _write_pivot_text(pivot_text: str, output_path: str) -> None:
    """Écrit le texte pivot vers le format de destination déterminé par l'extension."""
    ext = os.path.splitext(output_path)[1].lower()

    if ext in (".txt", ".md"):
        with open(output_path, "w", encoding="utf-8") as f:
            f.write(pivot_text)
        return

    if ext == ".html":
        # nl2br : sans ça, Markdown fusionne les lignes qui ne sont PAS séparées par une
        # ligne vide en un seul bloc de texte collé (cas très courant pour un simple .txt
        # ligne par ligne, ex. une liste) -> perte apparente du texte à l'affichage.
        body_html = markdown.markdown(pivot_text, extensions=["nl2br"])
        html_doc = (
            "<!DOCTYPE html>\n"
            "<html>\n<head>\n<meta charset=\"utf-8\">\n</head>\n<body>\n"
            f"{body_html}\n"
            "</body>\n</html>\n"
        )
        with open(output_path, "w", encoding="utf-8") as f:
            f.write(html_doc)
        return

    if ext == ".docx":
        document = docx.Document()
        for line in pivot_text.splitlines():
            if line.strip():
                document.add_paragraph(line)
        document.save(output_path)
        return

    if ext == ".pdf":
        _write_pdf(pivot_text, output_path)
        return

    raise ValueError(t("text.unsupported_target", ext=ext, accepted=", ".join(SUPPORTED_EXTENSIONS)))


def convert_text_file(input_path: str, output_path: str) -> None:
    """Détecte le format source et destination via l'extension de input_path/output_path
    et convertit. Lève une exception explicite si une extension n'est pas supportée."""
    in_ext = os.path.splitext(input_path)[1].lower()
    out_ext = os.path.splitext(output_path)[1].lower()

    if in_ext not in READABLE_EXTENSIONS:
        raise ValueError(t("text.unsupported_source", ext=in_ext, accepted=", ".join(READABLE_EXTENSIONS)))
    if out_ext not in SUPPORTED_EXTENSIONS:
        raise ValueError(t("text.unsupported_target", ext=out_ext, accepted=", ".join(SUPPORTED_EXTENSIONS)))

    out_dir = os.path.dirname(os.path.abspath(output_path))
    if out_dir:
        ensure_dir(out_dir)

    # Cas particulier : .doc -> .docx/.pdf via l'enregistrement natif de Word, qui
    # préserve la mise en page (voir _convert_doc_via_word). Pour toutes les autres
    # combinaisons, on passe par le pivot texte brut générique ci-dessous.
    if in_ext == ".doc" and out_ext in _WORD_SAVE_FORMAT:
        _convert_doc_via_word(input_path, output_path, out_ext)
        return

    pivot_text = _read_pivot_text(input_path)
    _write_pivot_text(pivot_text, output_path)


def build_output_path(input_path: str, output_dir: str, target_ext: str) -> str:
    """Construit le chemin de sortie : même nom que l'entrée, nouvelle extension, dans output_dir."""
    base_name = os.path.splitext(os.path.basename(input_path))[0]
    if not target_ext.startswith("."):
        target_ext = "." + target_ext
    return os.path.join(output_dir, base_name + target_ext)


# ------------------------------------------------------------------------------- UI ---

class TextConverterTab(ctk.CTkFrame):
    def __init__(self, master, **kwargs):
        super().__init__(master, fg_color="transparent", **kwargs)

        self.selected_files = []
        self.log_queue = queue.Queue()
        self.convert_thread = None

        self._build_ui()
        self.after(100, self._poll_queue)

    # ---------------------------------------------------------------- UI ---
    def _build_ui(self):
        body = build_scrollable_body(self)

        files_frame = ctk.CTkFrame(body)
        files_frame.pack(fill="x", **PAD)
        ctk.CTkLabel(files_frame, text=t("text.files_label"),
                     anchor="w").pack(fill="x", padx=8, pady=(8, 0))
        row = ctk.CTkFrame(files_frame, fg_color="transparent")
        row.pack(fill="x", padx=8, pady=8)
        self.files_entry = ctk.CTkEntry(row, placeholder_text=t("text.files_placeholder"))
        self.files_entry.pack(side="left", fill="x", expand=True)
        ctk.CTkButton(row, text=t("common.browse"), width=100, command=self._choose_files).pack(side="left", padx=(8, 0))

        enable_file_drop(files_frame, self._set_selected_files, extensions=set(READABLE_EXTENSIONS))

        opts_frame = ctk.CTkFrame(body)
        opts_frame.pack(fill="x", **PAD)
        fmt_col = ctk.CTkFrame(opts_frame, fg_color="transparent")
        fmt_col.pack(side="left", padx=8, pady=8)
        ctk.CTkLabel(fmt_col, text=t("text.convert_to")).pack(anchor="w")
        self.target_var = ctk.StringVar(value=".pdf")
        self.target_menu = ctk.CTkOptionMenu(fmt_col, values=[".txt", ".md", ".html", ".docx", ".pdf"],
                                              variable=self.target_var)
        self.target_menu.pack(anchor="w", pady=4)

        out_frame = ctk.CTkFrame(body)
        out_frame.pack(fill="x", **PAD)
        ctk.CTkLabel(out_frame, text=t("common.destination_folder"), anchor="w").pack(fill="x", padx=8, pady=(8, 0))
        out_row = ctk.CTkFrame(out_frame, fg_color="transparent")
        out_row.pack(fill="x", padx=8, pady=8)
        self.output_entry = ctk.CTkEntry(out_row)
        self.output_entry.insert(0, DEFAULT_CONVERTER_OUTPUT_DIR)
        self.output_entry.pack(side="left", fill="x", expand=True)
        ctk.CTkButton(out_row, text=t("common.browse"), width=100, command=self._choose_folder).pack(side="left", padx=(8, 0))

        btn_frame = ctk.CTkFrame(body, fg_color="transparent")
        btn_frame.pack(fill="x", **PAD)
        self.convert_btn = ctk.CTkButton(btn_frame, text=t("common.convert"), command=self._start_conversion)
        self.convert_btn.pack(side="left")
        ctk.CTkButton(btn_frame, text=t("common.open_folder"), command=self._open_folder).pack(side="right")
        ctk.CTkButton(btn_frame, text=t("common.export_logs"), width=140,
                      command=lambda: export_log_to_file(self.log_box, t("logname.text"))).pack(side="right", padx=(0, 8))

        self.progress_bar = ctk.CTkProgressBar(body)
        self.progress_bar.set(0)
        self.progress_bar.pack(fill="x", **PAD)

        self.status_label = ctk.CTkLabel(body, text=t("common.ready"), anchor="w")
        self.status_label.pack(fill="x", padx=16)

        self.log_box = ctk.CTkTextbox(body, height=200)
        self.log_box.pack(fill="both", expand=True, **PAD)
        self.log_box.configure(state="disabled")

    def _choose_files(self):
        files = filedialog.askopenfilenames(
            title=t("text.choose_files_title"),
            filetypes=[
                (t("text.filetype_documents"), "*.txt *.md *.html *.htm *.docx *.doc *.pdf"),
                (t("common.all_files"), "*.*"),
            ],
        )
        if files:
            self._set_selected_files(list(files))

    def _set_selected_files(self, files):
        self.selected_files = list(files)
        self.files_entry.delete(0, "end")
        if len(self.selected_files) == 1:
            self.files_entry.insert(0, self.selected_files[0])
        else:
            self.files_entry.insert(0, t("text.n_files_selected", n=len(self.selected_files)))

    def _choose_folder(self):
        folder = filedialog.askdirectory(initialdir=self.output_entry.get() or DEFAULT_CONVERTER_OUTPUT_DIR)
        if folder:
            self.output_entry.delete(0, "end")
            self.output_entry.insert(0, folder)

    def _open_folder(self):
        folder = ensure_dir(self.output_entry.get().strip() or DEFAULT_CONVERTER_OUTPUT_DIR)
        os.startfile(folder)

    # ------------------------------------------------------------ Logging ---
    def _log(self, message):
        self.log_queue.put(("log", message))

    def _poll_queue(self):
        try:
            while True:
                kind, payload = self.log_queue.get_nowait()
                if kind == "log":
                    self.log_box.configure(state="normal")
                    self.log_box.insert("end", payload + "\n")
                    self.log_box.see("end")
                    self.log_box.configure(state="disabled")
                elif kind == "progress":
                    self.progress_bar.set(payload)
                elif kind == "status":
                    self.status_label.configure(text=payload)
                elif kind == "done":
                    self.convert_btn.configure(state="normal")
        except queue.Empty:
            pass
        self.after(100, self._poll_queue)

    # -------------------------------------------------------- Conversion ---
    def _start_conversion(self):
        if not self.selected_files:
            messagebox.showwarning(t("common.no_file_title"), t("text.no_file_message"))
            return

        output_dir = ensure_dir(self.output_entry.get().strip() or DEFAULT_CONVERTER_OUTPUT_DIR)
        target_ext = self.target_var.get()
        files = list(self.selected_files)

        self.convert_btn.configure(state="disabled")
        self.progress_bar.set(0)
        self.status_label.configure(text=t("common.preparing"))

        self.convert_thread = threading.Thread(
            target=self._run_conversion, args=(files, target_ext, output_dir), daemon=True)
        self.convert_thread.start()

    def _run_conversion(self, files, target_ext, output_dir):
        total = len(files)
        try:
            for i, input_path in enumerate(files, start=1):
                filename = os.path.basename(input_path)
                self.log_queue.put(("status", t("text.converting", name=filename, index=i, total=total)))
                try:
                    output_path = build_output_path(input_path, output_dir, target_ext)
                    convert_text_file(input_path, output_path)
                    self._log(t("text.file_ok", name=filename, output=os.path.basename(output_path)))
                except Exception as e:
                    self._log(t("text.file_failed", name=filename, error=e))
                self.log_queue.put(("progress", i / total))

            self.log_queue.put(("status", t("common.done")))
            self._log(t("text.finished"))
        except Exception:
            self.log_queue.put(("status", t("common.unexpected_error_status")))
            self._log(t("common.unexpected_error_log") + "\n" + traceback.format_exc())
        finally:
            self.log_queue.put(("done", None))
